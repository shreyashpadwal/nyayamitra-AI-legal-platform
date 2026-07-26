import json
import io
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from datetime import datetime
from docx import Document

from ...db.database import get_db
from ...db.models import User, Conversation, Draft
from ...core.auth import get_current_user
from ...schemas.schemas import QuestionRequest, DocumentRequest
from ...services.vector_service import vector_service
from ...services.citizen_graph import stream_citizen_pipeline

router = APIRouter(tags=["Citizen - RAG & Drafting"])

def text_to_docx(text: str, title: str):
    doc = Document()
    doc.add_heading(title, 0)
    for line in text.split('\n'):
        doc.add_paragraph(line)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

@router.get("/download-docx/{draft_id}")
def download_docx(draft_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    draft = db.query(Draft).filter(Draft.id == draft_id, Draft.user_id == current_user.id).first()
    if not draft: raise HTTPException(status_code=404, detail="Not found")
    
    file_stream = text_to_docx(draft.content, draft.title)
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={draft.title.replace(' ', '_')}.docx"}
    )

@router.post("/ask")
def ask_question(body: QuestionRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    if not body.question.strip():
        raise HTTPException(status_code=400, detail="Empty question")

    result = vector_service.get_citizen_answer(
        body.question, "general", "Provide a helpful legal answer.",
        history=body.history or []
    )

    db.add(Conversation(user_id=current_user.id, question=body.question, answer=result["answer"]))
    db.commit()
    return result


@router.post("/ask-stream")
def ask_question_stream(
    body: QuestionRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Streaming SSE endpoint — same as /ask but yields incremental status events
    so the frontend can show live pipeline progress.

    SSE event format:
      data: {"type": "status", "message": "Searching the legal database..."}
      data: {"type": "final",  "answer": "...", "sources": [...]}
      data: {"type": "error",  "message": "..."}

    The original /citizen/ask endpoint is completely unchanged as a fallback.
    """
    if not body.question.strip():
        raise HTTPException(status_code=400, detail="Empty question")

    question = body.question

    def _event_generator():
        final_answer = ""
        try:
            for event_str in stream_citizen_pipeline(question, history=body.history or []):
                yield event_str
                # Capture answer from final event so we can persist it to DB
                if '"type": "final"' in event_str or '"type":"final"' in event_str:
                    try:
                        data_part = event_str.strip()
                        if data_part.startswith("data: "):
                            payload = json.loads(data_part[6:])
                            final_answer = payload.get("answer", "")
                    except Exception:
                        pass
        finally:
            if final_answer:
                try:
                    db.add(Conversation(
                        user_id=current_user.id,
                        question=question,
                        answer=final_answer,
                    ))
                    db.commit()
                except Exception:
                    db.rollback()

    return StreamingResponse(
        _event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )

@router.post("/generate-document")
def generate_doc(body: DocumentRequest, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    desc = body.fields.get("description", "")
    result = vector_service.generate_legal_document(body.doc_type, desc)

    # generate_legal_document now returns {"content": str, "verification": {...}}
    content = result["content"]
    verification = result["verification"]

    draft = Draft(
        user_id=current_user.id,
        doc_type=body.doc_type,
        title=f"{body.doc_type.replace('_', ' ').title()} - {datetime.now().strftime('%d %b %Y')}",
        content=content,
        form_data=json.dumps(body.fields)
    )
    db.add(draft)
    db.commit()
    return {
        "content": content,
        "title": draft.title,
        "draft_id": draft.id,
        "verification": verification,
    }

@router.get("/history")
def get_history(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return (
        db.query(Conversation)
        .filter(Conversation.user_id == current_user.id)
        .order_by(Conversation.timestamp.desc())
        .all()
    )

@router.delete("/history/{item_id}")
def delete_history(item_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    item = db.query(Conversation).filter(Conversation.id == item_id, Conversation.user_id == current_user.id).first()
    if not item: raise HTTPException(status_code=404, detail="Not found")
    db.delete(item)
    db.commit()
    return {"message": "Deleted"}

@router.get("/drafts")
def get_drafts(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    return db.query(Draft).filter(Draft.user_id == current_user.id).all()

@router.get("/drafts/{draft_id}")
def get_draft(draft_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    draft = db.query(Draft).filter(Draft.id == draft_id, Draft.user_id == current_user.id).first()
    if not draft: raise HTTPException(status_code=404, detail="Not found")
    return draft

@router.delete("/drafts/{draft_id}")
def delete_draft(draft_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    draft = db.query(Draft).filter(Draft.id == draft_id, Draft.user_id == current_user.id).first()
    if not draft: raise HTTPException(status_code=404, detail="Not found")
    db.delete(draft)
    db.commit()
    return {"message": "Deleted"}
